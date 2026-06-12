import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = "PatriotGroupCopierToner.md"
OUT_FILE = "PatriotGroupCopierToner.docx"


def add_inline_runs(para, text):
    """Parse inline **bold** and add runs to paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def parse_md_to_docx(md_path, out_path):
    doc = Document()

    # Narrow margins for readability
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Inline image
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            img_path = os.path.join(os.path.dirname(md_path), m.group(2))
            if os.path.isfile(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f"[Image not found: {m.group(2)}]")
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 9))
            heading.clear()
            add_inline_runs(heading, text)
            i += 1
            continue

        # Fenced code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph()
            para.style = "No Spacing"
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F5F5")
            para._p.get_or_add_pPr().append(shd)
            continue

        # Table — collect all consecutive table lines
        if re.match(r"^\|", line.strip()):
            table_lines = []
            while i < len(lines) and re.match(r"^\|", lines[i].strip()):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            # Filter out separator rows (---|---)
            rows = [r for r in table_lines if not re.match(r"^\|\s*[-:]+[\s|:-]*$", r)]
            if not rows:
                continue
            cols = [c.strip() for c in rows[0].strip("|").split("|")]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            # Header row
            hdr = table.rows[0].cells
            for ci, col in enumerate(cols):
                hdr[ci].text = ""
                p = hdr[ci].paragraphs[0]
                add_inline_runs(p, col)
                for run in p.runs:
                    run.bold = True
                set_cell_shading(hdr[ci], "D9E1F2")
            # Data rows
            for row_text in rows[1:]:
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                row = table.add_row().cells
                for ci, cell_text in enumerate(cells):
                    if ci < len(row):
                        row[ci].text = ""
                        add_inline_runs(row[ci].paragraphs[0], cell_text)
            doc.add_paragraph()  # spacing after table
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            style = "List Bullet 2" if indent > 0 else "List Bullet"
            para = doc.add_paragraph(style=style)
            para.clear()
            add_inline_runs(para, text)
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            style = "List Number 2" if indent > 0 else "List Number"
            para = doc.add_paragraph(style=style)
            para.clear()
            add_inline_runs(para, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        add_inline_runs(para, line.strip())
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    base = os.path.dirname(__file__)
    parse_md_to_docx(
        os.path.join(base, MD_FILE),
        os.path.join(base, OUT_FILE),
    )
