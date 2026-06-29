import argparse
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_inline_runs(para, text):
    """Parse inline **bold** and add runs to paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def parse_md_to_docx(md_path, out_path):
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    beside_image_path = None

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # <!-- beside: path --> directive — image placed to the right of the next table
        m_beside = re.match(r"<!--\s*beside:\s*(.+?)\s*-->", line.strip())
        if m_beside:
            beside_image_path = m_beside.group(1).strip()
            i += 1
            continue

        # Inline image (standalone line)
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), m.group(2)))
            if os.path.isfile(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f"[Image not found: {m.group(2)}]")
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 9))
            heading.clear()
            add_inline_runs(heading, text)
            i += 1
            continue

        # Fenced code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph()
            para.style = "No Spacing"
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F5F5")
            para._p.get_or_add_pPr().append(shd)
            continue

        # Table
        if re.match(r"^\|", line.strip()):
            table_lines = []
            while i < len(lines) and re.match(r"^\|", lines[i].strip()):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            rows = [r for r in table_lines if not re.match(r"^\|\s*[-:]+[\s|:-]*$", r)]
            if not rows:
                continue
            cols = [c.strip() for c in rows[0].strip("|").split("|")]

            num_cols = len(cols) + (1 if beside_image_path else 0)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = "Table Grid"

            # Header row
            hdr = table.rows[0].cells
            for ci, col in enumerate(cols):
                hdr[ci].text = ""
                p = hdr[ci].paragraphs[0]
                add_inline_runs(p, col)
                for run in p.runs:
                    run.bold = True
                set_cell_shading(hdr[ci], "D9E1F2")
            if beside_image_path:
                # Image column header — no shading, no border
                set_cell_shading(hdr[len(cols)], "FFFFFF")
                remove_cell_borders(hdr[len(cols)])

            # Data rows
            for row_text in rows[1:]:
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                row = table.add_row().cells
                for ci, cell_text in enumerate(cells):
                    if ci < len(cols):
                        row[ci].text = ""
                        img_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", cell_text.strip())
                        if img_m:
                            img_abs = os.path.normpath(
                                os.path.join(os.path.dirname(md_path), img_m.group(2))
                            )
                            if os.path.isfile(img_abs):
                                para = row[ci].paragraphs[0]
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                para.add_run().add_picture(img_abs, width=Inches(1.5))
                        else:
                            add_inline_runs(row[ci].paragraphs[0], cell_text)
                if beside_image_path:
                    set_cell_shading(row[len(cols)], "FFFFFF")
                    remove_cell_borders(row[len(cols)])

            # Merge image column and insert image
            if beside_image_path:
                img_col = len(cols)
                total_rows = len(rows)  # header + data rows
                top_cell = table.cell(0, img_col)
                bottom_cell = table.cell(total_rows - 1, img_col)
                merged = top_cell.merge(bottom_cell)

                img_abs = os.path.normpath(
                    os.path.join(os.path.dirname(md_path), beside_image_path)
                )
                if os.path.isfile(img_abs):
                    para = merged.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(img_abs, width=Inches(0.75))
                else:
                    merged.paragraphs[0].add_run(f"[Image not found: {beside_image_path}]")

                beside_image_path = None

            doc.add_paragraph()
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            style = "List Bullet 2" if indent > 0 else "List Bullet"
            para = doc.add_paragraph(style=style)
            para.clear()
            add_inline_runs(para, text)
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            style = "List Number 2" if indent > 0 else "List Number"
            para = doc.add_paragraph(style=style)
            para.clear()
            add_inline_runs(para, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        add_inline_runs(para, line.strip())
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert a Markdown file to a Word document.")
    parser.add_argument("--input", required=True, help="Path to the source .md file")
    parser.add_argument("--output", required=True, help="Path for the output .docx file")
    args = parser.parse_args()
    parse_md_to_docx(args.input, args.output)
